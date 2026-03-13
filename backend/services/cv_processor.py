"""
TalentLens — CV Processor Service
Extracts raw text from PDF and DOCX files.
"""

import os

import fitz  # PyMuPDF
from docx import Document


def extract_text(file_path: str) -> str:
    """
    Extract text content from a CV file.
    Supports PDF and DOCX formats.
    """
    ext = os.path.splitext(file_path)[1].lower()

    if ext == ".pdf":
        return _extract_pdf(file_path)
    elif ext == ".docx":
        return _extract_docx(file_path)
    else:
        raise ValueError(f"Unsupported file format: {ext}")


def _extract_pdf(path: str) -> str:
    """Extract text from a PDF file using PyMuPDF."""
    doc = fitz.open(path)
    text_parts = []
    for page in doc:
        text_parts.append(page.get_text("text"))
    doc.close()
    return "\n".join(text_parts).strip()


def _extract_docx(path: str) -> str:
    """Extract text from a DOCX file."""
    doc = Document(path)
    text_parts = []

    for para in doc.paragraphs:
        if para.text.strip():
            text_parts.append(para.text)

    # Also extract text from tables
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                text_parts.append(row_text)

    return "\n".join(text_parts).strip()
